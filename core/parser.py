"""
Document parser — converts .docx and .pdf files to structured category blocks.
Strong section splitter for English lesson files.
"""

import os
import re
import logging
from docx import Document
from core.config import CATEGORY_KEYWORDS

logger = logging.getLogger(__name__)


SECTION_PATTERNS = {
    "links": re.compile(
        r"^(links?|url|website|web\s*site|resources?|https?://|www\.)\b",
        re.I,
    ),
    "visuals": re.compile(
        r"^(visuals?|image|picture|photo|diagram|chart|table|video|watch|look\s+at)\b",
        re.I,
    ),
    "vocabulary": re.compile(
        r"^(part\s*2\b|task\s*2\b|vocabulary|vocab|words?|definitions?|new\s*words?|"
        r"word\s*list|key\s*words?|glossary|useful\s*words?|match\s+the\s+words?|"
        r"word\s*bank|terms?|phrases?)\b",
        re.I,
    ),
    "speaking": re.compile(
        r"^(part\s*1\b|task\s*1\b|speaking|discussion|discuss|pair\s*work|group\s*work|"
        r"role\s*play|debate|presentation|students\s*discuss|real[-\s]*life\s*problem\s*solving)\b",
        re.I,
    ),
    "listening": re.compile(
        r"^(part\s*4\b|task\s*4\b|listening|listening\s*task|audio|track|recording|"
        r"listen\s+and|listen\s+to|before\s*listening|after\s*listening|fact\s*hunt)\b",
        re.I,
    ),
    "reading": re.compile(
        r"^(part\s*3\b|task\s*3\b|reading|reading\s*task|reading\s*passage|reading\s*text|"
        r"read\s+the\s+text|read\s+the\s+following|article|passage|text\s*[a-z]?|"
        r"data\s*processing\s*day)\b",
        re.I,
    ),
    "writing": re.compile(
        r"^(writing|writing\s*task|write|essay|paragraph|composition|critical\s*thinking|"
        r"reflection|summary)\b",
        re.I,
    ),
    "games": re.compile(
        r"^(game|games|puzzle|crossword|bingo|fun\s*activity|scramble|quizlet)\b",
        re.I,
    ),
    "homework": re.compile(
        r"^(homework|home\s*task|assignment|at\s*home|self[-\s]*study)\b",
        re.I,
    ),
    "test_quiz": re.compile(
        r"^(answer\s*key|answers?|answer\s*sheet|reading\s*answer\s*sheet|"
        r"listening\s*answer\s*sheet|test|quiz|exam|true\s*/?\s*false|"
        r"multiple\s*choice|choose\s+the\s+correct|final\s*task)\b",
        re.I,
    ),
}


ORDERED_CATEGORIES = [
    "links",
    "visuals",
    "vocabulary",
    "speaking",
    "listening",
    "reading",
    "writing",
    "games",
    "homework",
    "test_quiz",
]


def _normalize(text: str) -> str:
    text = text or ""
    text = text.replace("\xa0", " ")
    text = text.replace("–", "-").replace("—", "-")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _clean_title(text: str) -> str:
    text = _normalize(text)
    text = re.sub(r"^[#\-\*\•·▪▫●○\d\.\)\(\s]+", "", text)
    text = text.strip(":：-–— ").strip()
    return text


def _table_to_text(table) -> str:
    rows = []

    for row in table.rows:
        cells = []

        for cell in row.cells:
            cell_text = _normalize(cell.text)
            if cell_text:
                cells.append(cell_text)

        if cells:
            rows.append(" | ".join(cells))

    return "\n".join(rows)


def _smart_task_classifier(text: str) -> str | None:
    low = _clean_title(text).lower()

    if not low:
        return None

    if "answer key" in low or "answer sheet" in low:
        return "test_quiz"

    if "true / false" in low or "true/false" in low:
        return "test_quiz"

    if "multiple choice" in low or "choose the correct" in low:
        return "test_quiz"

    if low.startswith("part 1") or low.startswith("task 1"):
        if "vocabulary" in low or "word" in low:
            return "vocabulary"
        return "speaking"

    if low.startswith("part 2") or low.startswith("task 2"):
        if "listening" in low:
            return "listening"
        if "reading" in low:
            return "reading"
        return "vocabulary"

    if low.startswith("part 3") or low.startswith("task 3"):
        if "speaking" in low or "problem solving" in low or "discussion" in low:
            return "speaking"
        return "reading"

    if low.startswith("part 4") or low.startswith("task 4"):
        if "critical thinking" in low or "writing" in low:
            return "writing"
        return "listening"

    if "real-life problem solving" in low or "students discuss" in low:
        return "speaking"

    if "critical thinking" in low:
        return "writing"

    if low in {"words", "word", "definitions", "definition"}:
        return "vocabulary"

    if low in {"audio", "recording", "track"}:
        return "listening"

    return None


def _classify_section_title(text: str) -> str | None:
    title = _clean_title(text)

    if not title:
        return None

    low = title.lower()

    if len(low) > 180:
        return None

    smart = _smart_task_classifier(low)

    if smart:
        return smart

    for cat in ORDERED_CATEGORIES:
        pattern = SECTION_PATTERNS.get(cat)
        if pattern and pattern.search(low):
            return cat

    for cat in ORDERED_CATEGORIES:
        for kw in CATEGORY_KEYWORDS.get(cat, []):
            kw_low = kw.lower().strip()

            if not kw_low:
                continue

            if low == kw_low or low.startswith(kw_low + ":"):
                return cat

    return None


def _looks_like_heading_line(text: str) -> bool:
    clean = _clean_title(text)

    if not clean:
        return False

    if _classify_section_title(clean):
        return True

    if len(clean) > 120:
        return False

    words = clean.split()

    if len(words) <= 10 and clean.endswith(":"):
        return True

    if len(words) <= 8 and clean.isupper():
        return True

    if re.match(r"^(task|part|activity|exercise)\s*\d+\.?", clean, re.I):
        return True

    if re.match(r"^\d+\.\s+[A-Z]", clean):
        return True

    return False


def _is_heading(para) -> bool:
    text = _normalize(para.text)

    if not text:
        return False

    if para.style and para.style.name and para.style.name.startswith("Heading"):
        return True

    if _looks_like_heading_line(text):
        return True

    runs = [r for r in para.runs if r.text.strip()]

    if runs and all(r.bold for r in runs) and len(text) <= 120:
        return True

    return False


def _extract_docx_blocks(path: str) -> tuple[list[tuple[str, str]], str]:
    doc = Document(path)
    items: list[tuple[str, str]] = []
    lesson_title = ""

    for elem in doc.element.body:
        tag = elem.tag.split("}")[-1]

        if tag == "tbl":
            tbl = next((t for t in doc.tables if t._element is elem), None)

            if tbl:
                table_text = _table_to_text(tbl)

                if table_text:
                    first_line = table_text.splitlines()[0]
                    detected = _classify_section_title(first_line)
                    kind = "heading" if detected else "text"
                    items.append((kind, table_text))

            continue

        para = next((p for p in doc.paragraphs if p._element is elem), None)

        if para is None:
            continue

        text = _normalize(para.text)

        if not text:
            continue

        detected = _classify_section_title(text)

        if not lesson_title and not detected and "lesson" in text.lower() and len(text) <= 120:
            lesson_title = _clean_title(text)
            continue

        kind = "heading" if _is_heading(para) or detected else "text"
        items.append((kind, text))

    return items, lesson_title


def _extract_pdf_text(path: str) -> str:
    try:
        import pdfplumber

        pages = []

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""

                if page_text.strip():
                    pages.append(page_text)

        return "\n".join(pages)

    except Exception as e:
        logger.warning("pdfplumber failed, fallback to pypdf: %s", e)

    try:
        from pypdf import PdfReader
    except Exception:
        from PyPDF2 import PdfReader

    reader = PdfReader(path)
    pages = []

    for page in reader.pages:
        page_text = page.extract_text() or ""

        if page_text.strip():
            pages.append(page_text)

    return "\n".join(pages)


def _extract_pdf_blocks(path: str) -> tuple[list[tuple[str, str]], str]:
    text = _extract_pdf_text(path)

    raw_lines = text.splitlines()
    lines = []

    for raw in raw_lines:
        line = _normalize(raw)

        if not line:
            continue

        if re.fullmatch(r"\d+", line):
            continue

        lines.append(line)

    items: list[tuple[str, str]] = []
    lesson_title = ""

    for line in lines:
        detected = _classify_section_title(line)

        if not lesson_title and not detected and "lesson" in line.lower() and len(line) <= 120:
            lesson_title = _clean_title(line)

        kind = "heading" if detected or _looks_like_heading_line(line) else "text"
        items.append((kind, line))

    return items, lesson_title


def _guess_category_from_content(text: str) -> str | None:
    low = text.lower()

    if re.search(r"https?://|www\.", low):
        return "links"

    if re.search(
        r"\b(customer|delivery|product|payment|definition|definitions|means|match the word|word bank|vocabulary)\b",
        low,
    ):
        return "vocabulary"

    if re.search(r"\blisten\b|\baudio\b|\btrack\b|\brecording\b", low):
        return "listening"

    if re.search(r"\bread\b|\bpassage\b|\barticle\b|\btext\b", low):
        return "reading"

    if re.search(r"\bdiscuss\b|\bspeak\b|\bpresentation\b|\bpair work\b|\bgroup work\b", low):
        return "speaking"

    if re.search(r"\bwrite\b|\bessay\b|\bparagraph\b|\bcomposition\b", low):
        return "writing"

    if re.search(r"\bhomework\b|\bassignment\b", low):
        return "homework"

    if re.search(r"\bquiz\b|\btest\b|\btrue\s*/?\s*false\b|\bmultiple choice\b", low):
        return "test_quiz"

    return None


def _split_items_to_categories(
    items: list[tuple[str, str]],
    lesson_title: str,
) -> tuple[dict, str]:
    result: dict[str, list[str]] = {cat: [] for cat in CATEGORY_KEYWORDS}

    current_cat = None
    buffer: list[str] = []

    def flush():
        nonlocal buffer, current_cat

        text = "\n".join(buffer).strip()

        if not text:
            buffer = []
            return

        target_cat = current_cat or _guess_category_from_content(text)

        if target_cat:
            result.setdefault(target_cat, []).append(text)

        buffer = []

    for kind, text in items:
        clean = _clean_title(text)
        detected = _classify_section_title(clean)

        if detected:
            flush()
            current_cat = detected
            buffer.append(f"**{clean}**")
            continue

        if kind == "heading":
            possible = _guess_category_from_content(clean)

            if possible and possible != current_cat:
                flush()
                current_cat = possible
                buffer.append(f"**{clean}**")
                continue

            if current_cat:
                buffer.append(f"**{clean}**")
            elif not lesson_title:
                lesson_title = clean[:80]

            continue

        if current_cat:
            buffer.append(text)
        else:
            possible = _guess_category_from_content(text)

            if possible:
                current_cat = possible
                buffer.append(text)
            elif not lesson_title:
                lesson_title = clean[:80]

    flush()

    result = {k: v for k, v in result.items() if v}

    # Fallback: agar parser hech qanday section topolmasa,
    # document matni yo‘qolib ketmasin.
    if not result:
        all_text = []

        for kind, text in items:
            clean = _normalize(text)

            if clean:
                all_text.append(clean)

        full_text = "\n".join(all_text).strip()

        if full_text:
            guessed = _guess_category_from_content(full_text) or "reading"
            result[guessed] = [full_text]

    links = []

    for blocks in result.values():
        for block in blocks:
            found = re.findall(r"https?://\S+|www\.\S+", block)
            links.extend([x.strip(".,;()[]{}<>") for x in found])

    if links:
        old_links = result.get("links", [])
        result["links"] = old_links + links

    return result, lesson_title or "Lesson"


def parse_document(path: str) -> tuple[dict, str]:
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        items, lesson_title = _extract_docx_blocks(path)
    elif ext == ".pdf":
        items, lesson_title = _extract_pdf_blocks(path)
    else:
        raise ValueError("Unsupported file type. Please upload .docx or .pdf")

    result, lesson_title = _split_items_to_categories(items, lesson_title)

    logger.info(
        "Parsed '%s': categories=%s",
        lesson_title,
        list(result.keys()),
    )

    return result, lesson_title
